from PyPDF2 import PdfReader
import docx

def extract_text(file):
    filename = file.name.lower()

    # PDF
    if filename.endswith(".pdf"):
        reader = PdfReader(file)
        return " ".join(
            [page.extract_text() for page in reader.pages if page.extract_text()]
        )

    # DOCX
    elif filename.endswith(".docx"):
        doc = docx.Document(file)
        return " ".join([p.text for p in doc.paragraphs])

    # TXT or any text file
    else:
        return file.read().decode("utf-8")

# ##Scanned pdf feature
# from PyPDF2 import PdfReader
# import docx
# import easyocr
# import pdfplumber
# import numpy as np

# # Initialize OCR once
# ocr_reader = easyocr.Reader(['en'], gpu=False)

# def extract_text(file):
#     filename = file.name.lower()
#     text = ""

#     # ================= PDF =================
#     if filename.endswith(".pdf"):
#         # 1️⃣ Try normal text extraction
#         reader = PdfReader(file)
#         for page in reader.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text += page_text + " "

#         # 2️⃣ OCR fallback for scanned PDF
#         if not text.strip():
#             file.seek(0)  # IMPORTANT
#             with pdfplumber.open(file) as pdf:
#                 for page in pdf.pages:
#                     page_image = page.to_image(resolution=300).original
#                     page_np = np.array(page_image)
#                     ocr_text = ocr_reader.readtext(page_np, detail=0)
#                     text += " ".join(ocr_text) + " "

#         return text.strip()

#     # ================= DOCX =================
#     elif filename.endswith(".docx"):
#         doc = docx.Document(file)
#         return " ".join(p.text for p in doc.paragraphs)

#     # ================= TXT =================
#     else:
#         return file.read().decode("utf-8")

